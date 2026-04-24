"""
Build pricing-page-updated.docx
Edstellar Corporate Training Pricing Page — Full Content Document
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

NAVY  = RGBColor(0x2D, 0x2F, 0x6B)
LIME  = RGBColor(0xC5, 0xE8, 0x26)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK  = RGBColor(0x1A, 0x1A, 0x1A)
GREY  = RGBColor(0x6B, 0x6D, 0x7B)
RED   = RGBColor(0xD0, 0x2B, 0x2B)
LIME_BG = RGBColor(0xF4, 0xF9, 0xD6)
NAVY_BG = RGBColor(0xE8, 0xE9, 0xF3)

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)

# ── Default style ─────────────────────────────────────────────────────────────
style = doc.styles['Normal']
font  = style.font
font.name = 'Calibri'
font.size = Pt(10.5)


# ─── Helper functions ─────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'),   val.get('val', 'single'))
            el.set(qn('w:sz'),    val.get('sz', '4'))
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), val.get('color', 'auto'))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def add_hr(doc, color='2D2F6B'):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    return p


def tag_bar(doc, module_name):
    """Navy tag bar showing the library module filename."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    set_cell_bg(cell, '2D2F6B')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f'  DESIGN MODULE: {module_name}')
    run.bold       = True
    run.font.color.rgb = WHITE
    run.font.size  = Pt(9)
    run.font.name  = 'Calibri'
    cell.height    = Cm(0.65)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def section_heading(doc, text, level=2):
    p = doc.add_heading(text, level=level)
    run = p.runs[0]
    run.font.color.rgb = NAVY
    run.font.name = 'Calibri'
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    return p


def label(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = GREY
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(2)
    return p


def body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix:
        r = p.add_run(bold_prefix + ' ')
        r.bold = True
        r.font.name = 'Calibri'
        r.font.size = Pt(10.5)
    r2 = p.add_run(text)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(10.5)
    p.paragraph_format.space_after  = Pt(5)
    p.paragraph_format.space_before = Pt(0)
    return p


def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix + ': ')
        r.bold = True
        r.font.name = 'Calibri'
        r.font.size = Pt(10.5)
    r2 = p.add_run(text)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(10.5)
    p.paragraph_format.space_after  = Pt(3)
    return p


def stat_note(doc, stats_list):
    """Inline stat pills as bold inline text."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    for i, (val, lbl) in enumerate(stats_list):
        r1 = p.add_run(val)
        r1.bold = True
        r1.font.size = Pt(13)
        r1.font.color.rgb = NAVY
        r1.font.name = 'Calibri'
        r2 = p.add_run(' ' + lbl)
        r2.font.size = Pt(10)
        r2.font.color.rgb = GREY
        r2.font.name = 'Calibri'
        if i < len(stats_list) - 1:
            r3 = p.add_run('   |   ')
            r3.font.size = Pt(10)
            r3.font.color.rgb = GREY
    return p


def note_box(doc, text, bg='F4F9D6'):
    table = doc.add_table(rows=1, cols=1)
    cell  = table.cell(0, 0)
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.font.name  = 'Calibri'
    r.font.size  = Pt(10)
    r.font.italic = True
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def divider(doc):
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    add_hr(doc, 'CCCCCC')
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ══════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('EDSTELLAR')
r.bold = True
r.font.size = Pt(11)
r.font.color.rgb = GREY
r.font.name = 'Calibri'

doc.add_paragraph()

p = doc.add_heading('Corporate Training Prices and Packages', 0)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.color.rgb = NAVY
    run.font.name = 'Calibri'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Pricing Page Content Document  |  Developer Reference')
r.font.size = Pt(11)
r.font.color.rgb = GREY
r.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Prepared: April 2026  |  URL: edstellar.com/corporate-training-pricing')
r.font.size = Pt(10)
r.font.color.rgb = GREY
r.font.name = 'Calibri'

doc.add_paragraph()
add_hr(doc, '2D2F6B')
doc.add_paragraph()

# Meta block
table = doc.add_table(rows=5, cols=2)
table.style = 'Table Grid'
headers = ['Title Tag', 'Meta Description', 'H1 (Exact)', 'Primary Keywords', 'Secondary Keywords']
values  = [
    'Corporate Training Prices and Packages 2026 | Edstellar',
    'Compare flexible corporate training prices and packages. Starter from 120 trainee slots. ISO-certified. 2,000+ courses. Get a custom quote today.',
    'Corporate Training Prices and Packages',
    'corporate training prices | corporate training packages',
    'corporate training fees | corporate training rates | corporate training costs | corporate training pricing | cost of corporate training | enterprise training program pricing | training charges | corporate training program cost | instructor-led training cost | onsite training pricing | virtual corporate training packages | corporate training pricing models | affordable training packages'
]
for i, (h, v) in enumerate(zip(headers, values)):
    row = table.rows[i]
    set_cell_bg(row.cells[0], 'E8E9F3')
    row.cells[0].paragraphs[0].add_run(h).bold = True
    row.cells[0].paragraphs[0].runs[0].font.name = 'Calibri'
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(9.5)
    row.cells[1].paragraphs[0].add_run(v)
    row.cells[1].paragraphs[0].runs[0].font.name = 'Calibri'
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(9.5)

doc.add_paragraph()

# Page flow overview
section_heading(doc, 'Page Section Flow', level=2)
sections_flow = [
    ('01', 'Hero Section',                     'hero-classic-split  [ADAPTED]'),
    ('02', 'Client Logo Strip',                'edstellar-logo-wall or strip  [EXISTING]'),
    ('03', 'Industry Pricing Benchmark',       'edstellar-stats-icon-cards  [ADAPTED]'),
    ('04', 'Pricing Tables',                   'CUSTOM pricing table module  [EXISTING + ENHANCED]'),
    ('05', 'Package Comparison Guide',         'edstellar-process-horizontal-timeline  [ADAPTED]'),
    ('06', 'ROI and Value Framing',            'edstellar-benefit-pillars-4col  [ADAPTED]'),
    ('07', 'Delivery Modes and Pricing',       'edstellar-three-pillars-cards  [ADAPTED]'),
    ('08', 'What Is Included vs Alternatives', 'CUSTOM comparison table  [NEW]'),
    ('09', 'Trainer Quality Assurance',        'edstellar-closed-loop-cycle  [ADAPTED]'),
    ('10', 'Dedicated Learning Services Mgr',  'edstellar-process-vertical-stepper  [ADAPTED]'),
    ('11', 'Client Outcomes',                  'edstellar-success-stories-dark-version-with-image'),
    ('12', 'Testimonials',                     'edstellar-testimonials-section-with-small-user-thumbnail'),
    ('13', 'Industries We Train',              'edstellar-industries-icon-grid'),
    ('14', 'Platform Capabilities',            'edstellar-services-grid-12col  [CONDENSED]'),
    ('15', 'Free Trial Risk Reversal CTA',     'edstellar-cta-banner-lime'),
    ('16', 'Trust and Credentials Strip',      'CUSTOM credentials strip  [NEW]'),
    ('17', 'Expanded FAQ',                     'edstellar-faq-section  [EXPANDED]'),
    ('18', 'Contact and Quote Form',           'edstellar-form-section  [ENHANCED]'),
    ('19', 'Footer',                           'edstellar-footer'),
]
t = doc.add_table(rows=len(sections_flow)+1, cols=3)
t.style = 'Table Grid'
for i, hd in enumerate(['#', 'Section Name', 'Design Module']):
    set_cell_bg(t.rows[0].cells[i], '2D2F6B')
    run = t.rows[0].cells[i].paragraphs[0].add_run(hd)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.name = 'Calibri'
    run.font.size = Pt(9.5)
for i, (num, name, mod) in enumerate(sections_flow):
    row = t.rows[i+1]
    if i % 2 == 0:
        for c in row.cells: set_cell_bg(c, 'F4F5F7')
    row.cells[0].paragraphs[0].add_run(num).font.name = 'Calibri'
    row.cells[1].paragraphs[0].add_run(name).font.name = 'Calibri'
    row.cells[2].paragraphs[0].add_run(mod).font.name = 'Calibri'
    for c in row.cells:
        c.paragraphs[0].runs[0].font.size = Pt(9.5)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 01 — HERO
# ══════════════════════════════════════════════════════════════════════════════
tag_bar(doc, 'hero-classic-split  [ADAPTED FOR PRICING PAGE]')
label(doc, 'Section 01 — Hero Banner')
section_heading(doc, 'Corporate Training Prices and Packages', level=2)

body(doc, 'Supercharge your upskilling initiatives with customized corporate training packages. Flexible plans for teams of every size, from focused skill sprints to enterprise-wide learning programs.')
body(doc, 'Transparent pricing. No hidden fees. Dedicated support from day one.')

label(doc, 'Trust Checkpoints (3 items)')
bullet(doc, 'Flexible corporate training packages for teams of all sizes')
bullet(doc, 'Transparent pricing with no hidden costs or surprise fees')
bullet(doc, 'Custom quotes available for enterprise teams with unlimited needs')

label(doc, 'CTA Buttons')
bullet(doc, 'Primary (lime): View Pricing Options')
bullet(doc, 'Secondary (outline): Talk to an Expert')

label(doc, 'Breadcrumb')
body(doc, 'Home > Corporate Training Pricing')

label(doc, 'Stat Strip (3 items, right column below image)')
stat_note(doc, [('500+', 'Courses and Programs'), ('60+', 'Global Enterprise Clients'), ('ISO', '9001:2015 and 27001:2022 Certified')])

divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 02 — CLIENT LOGO STRIP
# ══════════════════════════════════════════════════════════════════════════════
tag_bar(doc, 'edstellar-logo-wall  [EXISTING]')
label(doc, 'Section 02 — Client Logo Strip')
section_heading(doc, 'Globally Trusted by 60+ Leading Enterprises', level=2)

body(doc, 'Organizations across industries rely on Edstellar corporate training packages to upskill their teams at scale. From Fortune 500 companies to fast-growing enterprises, our training delivery has earned global trust.')

label(doc, 'Logo List (display as logo pills or image tiles)')
bullet(doc, 'Microsoft, Amazon, Intel, Adobe, MediaTek')
bullet(doc, 'Tata Chemicals, Johnson and Johnson, Godrej, Aditya Birla Group')
bullet(doc, 'VISA, ABB, Emerson, Autodesk, Cummins, Kelloggs')
bullet(doc, 'NIIT, TVS Electronics, UWorld, UBER + 30 more enterprise logos')

note_box(doc, 'Developer note: Use animated horizontal scroll strip or static 6-column grid. All logos in SVG or PNG white/greyscale version.')
divider(doc)


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 03 — INDUSTRY PRICING BENCHMARK
# ══════════════════════════════════════════════════════════════════════════════
tag_bar(doc, 'edstellar-stats-icon-cards  [ADAPTED]')
label(doc, 'Section 03 — Industry Pricing Benchmark Context  [NEW]')
section_heading(doc, 'How Much Does Corporate Training Cost? Industry Benchmarks', level=2)

body(doc, 'Understanding corporate training fees starts with knowing what the market pays. According to the ATD 2024 State of the Industry Report, companies spent an average of $774 per learner on training in 2024. That figure varies significantly by organization size.')
body(doc, 'Small companies (under 999 employees) spend around $1,047 per employee. Mid-size organizations (1,000 to 9,999 employees) average $739 per learner. Large enterprises with over 10,000 employees bring that cost down to $398 per employee through scale.')
body(doc, 'Edstellar corporate training packages are structured to deliver significant savings against these benchmarks, especially for organizations committing to annual or enterprise-wide programs.')

label(doc, '4-Column Stat Cards (icon + number + label)')
t = doc.add_table(rows=2, cols=4)
t.style = 'Table Grid'
stats = [
    ('$774', 'Average spend per learner globally (ATD 2024)'),
    ('Up to 60%', 'Cost savings with Edstellar packages vs ad-hoc training vendors'),
    ('2,000+', 'Instructor-led corporate training programs available'),
    ('1,200+', 'Verified expert trainers across 100+ domains'),
]
for i, (val, lbl) in enumerate(stats):
    set_cell_bg(t.rows[0].cells[i], '2D2F6B')
    r = t.rows[0].cells[i].paragraphs[0].add_run(val)
    r.bold = True; r.font.color.rgb = LIME; r.font.size = Pt(14); r.font.name = 'Calibri'
    set_cell_bg(t.rows[1].cells[i], 'F4F5F7')
    r2 = t.rows[1].cells[i].paragraphs[0].add_run(lbl)
    r2.font.size = Pt(9); r2.font.name = 'Calibri'

doc.add_paragraph()

label(doc, 'Body Content Below Stats')
body(doc, 'Corporate training rates also vary by delivery format. Instructor-led training (ILT) typically costs $3,000 to $7,000 per finished training hour when sourced from ad-hoc vendors. Virtual instructor-led training (VILT) runs 10 to 30 percent lower. With Edstellar packages, organizations pre-purchase trainee licenses in bulk, reducing per-session costs and eliminating unpredictable vendor charges.')
body(doc, 'Our pricing model is straightforward: one trainee license equals one 8-hour session for one participant. Packages are structured around your team size and annual training volume, giving you full control over your corporate training budget.')

note_box(doc, 'SEO note: This section directly answers "how much does corporate training cost", "corporate training fees", "corporate training rates", "cost of corporate training", and "average cost of instructor-led training" queries from Search Console data.')
divider(doc)


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 04 — PRICING TABLES
# ══════════════════════════════════════════════════════════════════════════════
tag_bar(doc, 'CUSTOM PRICING TABLE MODULE  [EXISTING + ENHANCED]')
label(doc, 'Section 04 — Corporate Training Packages and Pricing Tables  [ENHANCED]')
section_heading(doc, 'Customized Corporate Training Packages for Organizations Globally', level=2)

body(doc, 'Choose a corporate training package that fits your team size, training frequency, and validity period. All packages include access to 2,000+ instructor-led training programs, a Dedicated Learning Services Manager, and flexible delivery via virtual or onsite sessions.')
body(doc, 'One trainee license equals one 8-hour session for one participant. Half-day sessions (4 hours) consume half a license. Packages are priced for budget clarity and long-term training planning.')

# Package overview table
label(doc, 'Package Comparison Table')
t = doc.add_table(rows=6, cols=5)
t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.CENTER

headers_row = ['Feature', 'Starter', 'Growth', 'Enterprise', 'Custom']
bg_colors   = ['2D2F6B', '2D2F6B', '2D2F6B', '2D2F6B', '2D2F6B']
for i, h in enumerate(headers_row):
    set_cell_bg(t.rows[0].cells[i], '2D2F6B')
    r = t.rows[0].cells[i].paragraphs[0].add_run(h)
    r.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(10); r.font.name = 'Calibri'

rows_data = [
    ('Validity',          '6 months',   '9 months',   '12 months',         '12 months'),
    ('Trainee Slots',     '120',        '320',        '800',                'Unlimited'),
    ('Training Hours',    '64 hours',   '160 hours',  '400 hours',          'Unlimited'),
    ('Sessions',          'Up to 10',   'Up to 25',   'Up to 60',           'Unlimited'),
    ('Savings',           '',           'Save up to 10%', 'Save up to 20%', 'Cost per employee'),
]
highlights = [False, False, True, False, False]  # Growth is most popular

for i, (row_data, highlight) in enumerate(zip(rows_data, highlights)):
    row = t.rows[i+1]
    for j, val in enumerate(row_data):
        bg = 'EBF0F8' if highlight and j > 0 else ('F4F5F7' if i % 2 == 0 else 'FFFFFF')
        if j == 0: bg = 'E8E9F3'
        set_cell_bg(row.cells[j], bg)
        r = row.cells[j].paragraphs[0].add_run(val)
        if j == 0: r.bold = True
        r.font.name = 'Calibri'
        r.font.size = Pt(10)

doc.add_paragraph()

label(doc, 'Best-For Labels (shown below each package name in the table)')
bullet(doc, 'Starter: Organizations beginning their structured training journey or with focused annual training requirements')
bullet(doc, 'Growth: Organizations with moderate training needs across multiple teams or departments (MOST POPULAR badge)')
bullet(doc, 'Enterprise: Large organizations running continuous, multi-domain training programs across multiple locations')
bullet(doc, 'Custom: Enterprise teams requiring unlimited sessions, custom delivery formats, or specialized program structures')

label(doc, 'CTA per package')
body(doc, 'Button label: Enquire Now (links to contact form / quote request)')

label(doc, 'Common Features Across All Packages (checklist below table)')
bullet(doc, '2,000+ instructor-led training programs (IT, management, leadership, behavioral, compliance, AI, and more)')
bullet(doc, 'Virtual instructor-led training (VILT) and onsite delivery options')
bullet(doc, 'Dedicated Learning Services Manager as single point of contact')
bullet(doc, 'Retrospective sessions (2 hours per training program)')
bullet(doc, 'Paid trial session available before committing to a package')
bullet(doc, 'Access to add-on services including coaching, mentoring, TNA, and content development')

label(doc, 'Pricing Footnotes (displayed below the table)')
bullet(doc, '[1] One trainee license equals one 8-hour session for one participant.')
bullet(doc, '[2] Half-day training lasts 4 hours and uses half a trainee license slot.')
bullet(doc, '[3] Group size is flexible within the allotted trainee license count per package.')
bullet(doc, '[4] Onsite in-person training is charged at 30% extra to cover trainer travel and accommodation. Additional fees may apply for remote locations.')
bullet(doc, '[5] The Custom Package is priced on a cost-per-employee basis and includes unlimited sessions tailored to organizational needs.')

label(doc, 'What Is Not Included (small note below footnotes)')
bullet(doc, 'Certification and exam fees')
bullet(doc, 'Post-training support and follow-up assessments')
bullet(doc, 'Lab access and technical environment setup')
bullet(doc, 'Venue, facilities, and catering for onsite sessions')
bullet(doc, 'Certain specialized or advanced training sessions as noted at time of booking')

divider(doc)


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 05 — PACKAGE COMPARISON GUIDE
# ══════════════════════════════════════════════════════════════════════════════
tag_bar(doc, 'edstellar-process-horizontal-timeline  [ADAPTED]')
label(doc, 'Section 05 — Package Comparison Guide  [NEW]')
section_heading(doc, 'How to Choose the Right Corporate Training Package', level=2)

body(doc, 'Selecting the right corporate training package depends on three key factors: your team size, your annual training volume, and your preferred validity window. This guide helps you match your organization to the right plan.')

label(doc, '3-Step Selector (horizontal timeline / step flow)')
bullet(doc, 'Step 1 — Estimate Your Team Size: Count the total number of employees who will need training over the next 6 to 12 months. Include all departments, locations, and seniority levels in scope.')
bullet(doc, 'Step 2 — Calculate Training Hours Needed: Multiply your team size by the average training hours per employee per year. Industry average is 42 hours per employee annually (ATD). Adjust for your specific skill gap priorities.')
bullet(doc, 'Step 3 — Choose Your Validity Window: Match the package validity to your planning cycle. Use 6-month packages for focused sprints. Choose 9 or 12-month packages for annual workforce development programs.')

label(doc, 'Package Selector Reference (use as inline decision table or callout)')
t2 = doc.add_table(rows=5, cols=4)
t2.style = 'Table Grid'
t2_headers = ['Team Size', 'Annual Hours', 'Recommended Package', 'Common Use Case']
for i, h in enumerate(t2_headers):
    set_cell_bg(t2.rows[0].cells[i], '2D2F6B')
    r = t2.rows[0].cells[i].paragraphs[0].add_run(h)
    r.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(9.5); r.font.name = 'Calibri'
rows2 = [
    ('Up to 15 employees',     'Up to 64 hrs',  'Starter',    'New hire onboarding, focused skill sprint, compliance refresh'),
    ('16 to 40 employees',     '65 to 160 hrs', 'Growth',     'Multi-team upskilling, leadership cohort, annual L&D plan'),
    ('41 to 100 employees',    '161 to 400 hrs','Enterprise', 'Division-wide programs, multi-domain training, continuous delivery'),
    ('100+ employees',         'Unlimited',     'Custom',     'Enterprise rollout, per-employee pricing, tailored program design'),
]
for i, row_data in enumerate(rows2):
    row = t2.rows[i+1]
    for j, val in enumerate(row_data):
        bg = 'F4F5F7' if i % 2 == 0 else 'FFFFFF'
        set_cell_bg(row.cells[j], bg)
        r = row.cells[j].paragraphs[0].add_run(val)
        r.font.name = 'Calibri'; r.font.size = Pt(9.5)
        if j == 2: r.bold = True

doc.add_paragraph()
body(doc, 'Not sure which package fits your needs? Our team can build a custom training plan based on your specific workforce goals. Use the quote form at the bottom of this page to get a tailored recommendation.')

note_box(doc, 'SEO note: Addresses "enterprise training program pricing", "corporate training pricing models", "corporate training module pricing and multi-seat discounts for teams", and "yearly corporate training program packages providers" from Search Console.')
divider(doc)


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 06 — ROI AND VALUE FRAMING
# ══════════════════════════════════════════════════════════════════════════════
tag_bar(doc, 'edstellar-benefit-pillars-4col  [ADAPTED — navy background]')
label(doc, 'Section 06 — ROI and Value Framing  [NEW]')
section_heading(doc, 'Corporate Training Is an Investment, Not a Cost', level=2)

body(doc, 'Every line item in your corporate training budget has a return. Research consistently shows that organizations investing in structured employee training outperform those that do not, across revenue growth, retention, and innovation.')
body(doc, 'Here is what the data shows about the value of corporate training programs:')

label(doc, '4-Column Benefit Pillars (numbered, navy background, lime accents)')
bullet(doc, '01 — Reduce Employee Turnover Costs: Replacing a mid-level employee costs between 50 and 200 percent of their annual salary (SHRM). A structured training program reduces turnover by improving engagement, career visibility, and role confidence. The cost of your corporate training package is a fraction of one replacement hire.')
bullet(doc, '02 — Drive Revenue Through Skills: A 2 percent improvement in sales conversion from skills training generates $200,000 in additional revenue for every $10 million a sales team produces. Corporate training fees pay for themselves when tied to measurable performance outcomes.')
bullet(doc, '03 — Build a Future-Ready Workforce: Companies with strong L&D programs are 92 percent more likely to innovate and 46 percent more likely to be first to market (Deloitte Human Capital Trends). Investing in the right corporate training packages today reduces skills debt tomorrow.')
bullet(doc, '04 — Faster Time to Competency: Structured instructor-led training reduces time to competency by 30 to 50 percent compared to on-the-job learning alone (McKinsey). Teams trained through Edstellar packages apply skills within days of program completion.')

body(doc, 'When comparing corporate training costs across providers, factor in what you are actually buying: trainer expertise, program quality, management support, and measurable outcomes. Edstellar packages are designed to deliver all four.')

note_box(doc, 'SEO note: Addresses "corporate training costs", "cost of corporate training", "is corporate training worth it", "corporate training program cost" and broader ROI-intent queries.')
divider(doc)


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 07 — DELIVERY MODES AND PRICING
# ══════════════════════════════════════════════════════════════════════════════
tag_bar(doc, 'edstellar-three-pillars-cards  [ADAPTED]')
label(doc, 'Section 07 — Delivery Modes and Pricing  [NEW]')
section_heading(doc, 'Virtual, Onsite, and Blended Corporate Training Pricing Explained', level=2)

body(doc, 'Edstellar supports multiple delivery formats within every corporate training package. The pricing structure adapts to how and where your teams learn best.')

label(doc, '3-Column Card Layout (icon + title + description)')
bullet(doc, 'Virtual Instructor-Led Training (VILT): All packages include virtual delivery at no extra cost. Live, trainer-facilitated sessions conducted via your preferred video platform. Ideal for distributed teams across multiple offices or geographies. Typical corporate training rates for VILT from ad-hoc vendors range from $2,000 to $5,000 per session. With Edstellar packages, VILT is included within your trainee license allocation.')
bullet(doc, 'Onsite Instructor-Led Training: In-person training delivered at your office or preferred venue. Onsite sessions are charged at 30 percent above the standard package rate to cover trainer travel and accommodation costs. Particularly effective for hands-on technical training, leadership workshops, and compliance programs requiring face-to-face delivery. Additional fees apply for remote or international locations.')
bullet(doc, 'Blended Learning Add-On: Combine virtual instructor-led sessions with self-paced eLearning modules to extend learning beyond classroom hours. Available as an add-on to any package tier. Blended delivery improves knowledge retention and is ideal for technical training, certifications, and behavioral programs where reinforcement matters.')

label(doc, 'Supporting Note (below cards)')
body(doc, 'All delivery modes are coordinated by your Dedicated Learning Services Manager. They handle scheduling, trainer briefing, logistics, and post-session feedback collection across all format types.')

note_box(doc, 'SEO note: Directly targets "onsite corporate training", "cost of in-person corporate training", "how much does onsite training cost", "virtual corporate training packages", "in-person training packages" from Search Console.')
divider(doc)


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 08 — WHAT IS INCLUDED VS ALTERNATIVES
# ══════════════════════════════════════════════════════════════════════════════
tag_bar(doc, 'CUSTOM COMPARISON TABLE MODULE  [NEW]')
label(doc, 'Section 08 — What Is Included vs Alternatives  [NEW]')
section_heading(doc, 'Why Edstellar Corporate Training Packages Deliver More Value', level=2)

body(doc, 'When evaluating corporate training solutions, organizations typically compare four options: building an in-house training team, booking ad-hoc training vendors, using eLearning platforms, or partnering with a full-service training provider like Edstellar.')
body(doc, 'Here is how those options compare on the factors that matter most:')

label(doc, 'Comparison Table (4 columns, checkmarks and X marks)')
t3 = doc.add_table(rows=9, cols=5)
t3.style = 'Table Grid'
t3_headers = ['Feature', 'Edstellar Packages', 'In-House Team', 'Ad-Hoc Vendors', 'eLearning Platforms']
for i, h in enumerate(t3_headers):
    set_cell_bg(t3.rows[0].cells[i], '2D2F6B')
    r = t3.rows[0].cells[i].paragraphs[0].add_run(h)
    r.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(9); r.font.name = 'Calibri'

rows3 = [
    ('2,000+ course library',           'Yes', 'Limited',     'Varies',    'Yes (self-paced only)'),
    ('Live expert trainers on demand',  'Yes', 'Yes (costly)', 'Yes',      'No'),
    ('Dedicated program manager',       'Yes', 'Maybe',       'Rarely',    'No'),
    ('Flexible package pricing',        'Yes', 'High fixed cost', 'Per session only', 'Per seat/year'),
    ('Virtual and onsite delivery',     'Yes', 'Yes',         'Yes',       'Virtual only'),
    ('Training management platform',   'Yes', 'No',          'No',        'Yes (basic)'),
    ('Transparent corporate training fees', 'Yes', 'Hidden overhead', 'Quote-based', 'Published rates'),
    ('ISO-certified quality assurance', 'Yes', 'No',          'Varies',    'No'),
]
for i, row_data in enumerate(rows3):
    row = t3.rows[i+1]
    for j, val in enumerate(row_data):
        bg = 'F4F9D6' if j == 1 else ('F4F5F7' if i % 2 == 0 else 'FFFFFF')
        if j == 0: bg = 'E8E9F3'
        set_cell_bg(row.cells[j], bg)
        r = row.cells[j].paragraphs[0].add_run(val)
        r.font.name = 'Calibri'; r.font.size = Pt(9)
        if j == 0: r.bold = True
        if val == 'Yes': r.font.color.rgb = RGBColor(0x1A, 0x6E, 0x1A)

doc.add_paragraph()
body(doc, 'Edstellar combines the breadth of an eLearning platform with the quality of live expert trainers, the management support of a dedicated program team, and the pricing flexibility of a structured package model. That combination is why over 60 enterprise clients trust us with their annual training budgets.')

note_box(doc, 'SEO note: Targets "corporate training companies", "corporate training solutions", "best corporate training solutions", "compare corporate training packages", and "full-service corporate learning vendors" queries.')
divider(doc)


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 09 — TRAINER QUALITY ASSURANCE
# ══════════════════════════════════════════════════════════════════════════════
tag_bar(doc, 'edstellar-closed-loop-cycle  [ADAPTED — 4-step vetting process]')
label(doc, 'Section 09 — Trainer Quality Assurance  [NEW]')
section_heading(doc, 'Every Corporate Training Program Is Delivered by Verified Expert Trainers', level=2)

body(doc, 'The value of any corporate training package ultimately comes down to trainer quality. Edstellar maintains a network of 1,200+ verified domain expert trainers, each cleared through a rigorous 4-step vetting process before delivering a single session.')

stat_note(doc, [
    ('1,200+', 'Verified trainers'),
    ('12 yrs', 'Average domain experience'),
    ('4.8/5', 'Average trainer rating'),
    ('100+', 'Skill domains covered')
])

label(doc, '4-Step Vetting Process (connected cycle or numbered steps)')
bullet(doc, 'Step 1 — Application Review: Trainers submit credentials, certifications, and a portfolio of past delivery. Industry experience, domain specialization, and organizational context are verified.')
bullet(doc, 'Step 2 — Subject Matter Assessment: Each trainer completes a domain knowledge assessment specific to their area of expertise. Technical trainers are tested against current standards and tool versions.')
bullet(doc, 'Step 3 — Live Delivery Audit: A sample training session is evaluated by Edstellar L&D specialists. Assessment criteria include clarity, engagement, pace, learner interaction, and content accuracy.')
bullet(doc, 'Step 4 — Client Feedback Loop: Post-session ratings from real clients feed back into trainer profiles. Trainers maintaining scores below 4.5/5 undergo re-certification or are removed from the active roster.')

body(doc, 'Trainer specializations span IT and technical domains (cloud, cybersecurity, DevOps, AI, data science), management and leadership, behavioral skills, compliance, and social impact training. When you purchase a corporate training package, you get access to the entire verified trainer pool across all these categories.')

note_box(doc, 'SEO note: Addresses "corporate trainer fees", "corporate trainer certification cost", "how much do corporate trainers charge", "cost of instructor-led team skills training providers USA" queries from Search Console.')
divider(doc)


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 10 — DEDICATED LEARNING SERVICES MANAGER
# ══════════════════════════════════════════════════════════════════════════════
tag_bar(doc, 'edstellar-process-vertical-stepper  [ADAPTED — navy bg, 5 steps]')
label(doc, 'Section 10 — Dedicated Learning Services Manager  [NEW]')
section_heading(doc, 'Every Package Comes with a Dedicated Learning Services Manager', level=2)

body(doc, 'One of the most significant differences between Edstellar corporate training packages and ad-hoc training vendors is the Dedicated Learning Services Manager (LSM) included in every plan. Your LSM is your single point of contact for the entire engagement, from onboarding to renewal.')
body(doc, 'This is what your LSM handles across the engagement lifecycle:')

label(doc, '5-Phase Lifecycle (vertical stepper, navy background, lime step numbers)')
bullet(doc, 'Phase 1 — Onboarding: The LSM aligns with your HR or L&D team to understand your training objectives, employee profiles, skill gap priorities, and scheduling preferences. A 12-month training calendar is mapped in week one.')
bullet(doc, 'Phase 2 — Scheduling and Coordination: The LSM matches program requirements with the right verified trainers, manages session booking, confirms participant lists, and sends pre-training briefing packs to both trainers and teams.')
bullet(doc, 'Phase 3 — Delivery Support: During live sessions, the LSM monitors attendance, manages logistics, and is available as an escalation point. For onsite sessions, the LSM coordinates all travel and venue logistics.')
bullet(doc, 'Phase 4 — Feedback and Retrospective: Within 48 hours of each session, the LSM collects trainer and participant feedback, facilitates a 2-hour retrospective, and documents improvement actions for the next session.')
bullet(doc, 'Phase 5 — Reporting and Renewal: Monthly utilization reports track trainee hours consumed, sessions completed, and completion rates. As your package validity approaches renewal, the LSM prepares a usage summary and recommends the right plan for the next cycle.')

body(doc, 'No other corporate training solution in this price range includes a fully managed service layer. Most training fees you pay elsewhere go to the trainer. With Edstellar, you also get the program infrastructure that makes training stick.')

divider(doc)


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 11 — CLIENT OUTCOMES
# ══════════════════════════════════════════════════════════════════════════════
tag_bar(doc, 'edstellar-success-stories-dark-version-with-image')
label(doc, 'Section 11 — Client Outcomes and Case Studies  [NEW]')
section_heading(doc, 'Real Results from Edstellar Corporate Training Packages', level=2)

body(doc, 'These are real outcomes from organizations that chose Edstellar training packages to address their workforce development goals.')

label(doc, 'Story 1 (IT and Technology Sector)')
body(doc, '', bold_prefix='Company Profile:')
body(doc, 'A global technology firm with 3,200 employees across six countries. Challenge: upskilling engineering teams on cloud architecture and cybersecurity within a 9-month window.')
body(doc, '', bold_prefix='Package Used:')
body(doc, 'Enterprise Package. 400 hours of training across 14 technical domains. Delivered via VILT for distributed teams.')
body(doc, '', bold_prefix='Outcome:')
body(doc, '840 engineers trained. 93 percent session completion rate. 31 percent reduction in cloud infrastructure incidents within 6 months of program completion.')
body(doc, '', bold_prefix='Client Quote:')
body(doc, '"Edstellar delivered a training program at a scale and pace we could not have managed with individual vendors. The LSM made the whole operation seamless." -- L&D Director, Global Technology Firm')

doc.add_paragraph()

label(doc, 'Story 2 (BFSI Sector)')
body(doc, '', bold_prefix='Company Profile:')
body(doc, 'A leading financial services company with 1,500 employees. Challenge: mandatory compliance and leadership training across regional offices in three countries.')
body(doc, '', bold_prefix='Package Used:')
body(doc, 'Growth Package. 160 hours of blended training covering regulatory compliance, risk management, and leadership development.')
body(doc, '', bold_prefix='Outcome:')
body(doc, '320 employees trained. 100 percent compliance training completion ahead of regulatory deadline. Leadership cohort reported 22 percent improvement in team engagement scores.')
body(doc, '', bold_prefix='Client Quote:')
body(doc, '"The structured package gave us cost predictability we have never had before with training. We knew exactly what we were spending and exactly what we were getting." -- VP of Human Resources, Financial Services Company')

doc.add_paragraph()

label(doc, 'Story 3 (Manufacturing Sector)')
body(doc, '', bold_prefix='Company Profile:')
body(doc, 'A mid-size manufacturing company with 450 employees across two plants. Challenge: upskilling operations and quality teams on lean manufacturing and safety compliance.')
body(doc, '', bold_prefix='Package Used:')
body(doc, 'Starter Package (later upgraded to Growth). 64 hours of onsite training delivered at plant locations.')
body(doc, '', bold_prefix='Outcome:')
body(doc, '120 employees trained in the first 6 months. Package upgraded at renewal. 18 percent reduction in quality defects reported post-training.')
body(doc, '', bold_prefix='Client Quote:')
body(doc, '"We started with the Starter package to test the quality. We were so impressed that we upgraded to Growth at renewal. The trainers Edstellar sent were genuinely expert." -- Operations Manager, Manufacturing Company')

note_box(doc, 'Developer note: Use 3-slide carousel with dark navy card design, image on left, quote and stats on right. Each card has company type, package tier used, and 2-3 key outcome stats with lime accent numbers.')
divider(doc)


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 12 — TESTIMONIALS
# ══════════════════════════════════════════════════════════════════════════════
tag_bar(doc, 'edstellar-testimonials-section-with-small-user-thumbnail')
label(doc, 'Section 12 — Client Testimonials  [NEW — pricing-specific]')
section_heading(doc, 'What L&D Leaders Say About Edstellar Training Packages', level=2)

body(doc, 'These testimonials come specifically from clients who evaluated multiple corporate training solutions before choosing Edstellar. Their feedback speaks to the value, quality, and support they experienced.')

label(doc, 'Testimonial 1')
body(doc, '"We compared three training providers on price and scope. Edstellar offered the most transparent corporate training pricing, the broadest course library, and the only real managed service. The Dedicated LSM alone justified the investment." -- Head of L&D, Technology Company, USA')

label(doc, 'Testimonial 2')
body(doc, '"The Growth package covered everything our team of 280 needed for the year. One package, one contact, one invoice. The simplicity of that versus managing 10 vendors is hard to overstate." -- HR Director, Pharma Company, India')

label(doc, 'Testimonial 3')
body(doc, '"Edstellar training fees are competitive but what you actually get is far beyond what the price suggests. The quality of trainers, the coordination, the follow-up. It does not feel like a vendor relationship. It feels like an extension of our L&D team." -- Senior Manager, Learning and Development, BFSI, UAE')

label(doc, 'Testimonial 4')
body(doc, '"We had budget pressure and needed affordable training packages that did not compromise on quality. The Starter package was exactly right for a team of our size and we saw results quickly." -- Training Coordinator, Retail Company, UK')

note_box(doc, 'Developer note: Display as split carousel with left heading + navigation arrows, right quote card with avatar initials circle, role, company type, and country.')
divider(doc)


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 13 — INDUSTRIES WE TRAIN
# ══════════════════════════════════════════════════════════════════════════════
tag_bar(doc, 'edstellar-industries-icon-grid')
label(doc, 'Section 13 — Industries We Train  [NEW]')
section_heading(doc, 'Corporate Training Packages Across Every Major Industry', level=2)

body(doc, 'Edstellar delivers corporate training programs across all major industry verticals. Our trainer network includes subject matter experts with direct industry experience, ensuring that training content is relevant to your specific business context.')

label(doc, '9-Industry Icon Grid (icon circle + industry name + 1-line description)')
t4 = doc.add_table(rows=5, cols=2)
t4.style = 'Table Grid'
industries = [
    ('IT and Technology',       'Cloud, cybersecurity, DevOps, AI, data science, and software development training'),
    ('BFSI',                    'Risk management, compliance, financial regulations, leadership, and digital banking'),
    ('Manufacturing',           'Lean manufacturing, quality management, safety compliance, and operations'),
    ('Healthcare and Pharma',   'Regulatory compliance, clinical operations, patient safety, and quality assurance'),
    ('Retail and E-Commerce',   'Customer service, sales excellence, supply chain, and digital retail operations'),
    ('Energy and Utilities',    'HSE compliance, technical skills, sustainability, and operations management'),
    ('Telecom',                 'Network operations, 5G technology, customer experience, and digital transformation'),
    ('Government and PSU',      'Leadership development, governance, compliance, and public sector digital skills'),
    ('Professional Services',   'Project management, consulting skills, communication, and client management'),
]
for i, (industry, desc) in enumerate(industries):
    row_idx = i // 2
    col_idx = i % 2
    if row_idx < len(t4.rows) and col_idx < 2:
        cell = t4.rows[row_idx].cells[col_idx]
        p = cell.paragraphs[0]
        r1 = p.add_run(industry + ': ')
        r1.bold = True; r1.font.name = 'Calibri'; r1.font.size = Pt(10)
        r2 = p.add_run(desc)
        r2.font.name = 'Calibri'; r2.font.size = Pt(10)

doc.add_paragraph()

body(doc, 'Each industry benefits from training programs that map to common frameworks and certifications within that sector. Our trainers understand the regulatory environment, the tooling, and the performance standards specific to your industry, making every session immediately applicable.')

note_box(doc, 'SEO note: Addresses "corporate training companies in USA", "customised corporate training worldwide", "corporate training in Dubai", "on-site corporate training India", "corporate IT training" and other geo and industry variants from Search Console data.')
divider(doc)


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 14 — PLATFORM CAPABILITIES
# ══════════════════════════════════════════════════════════════════════════════
tag_bar(doc, 'edstellar-services-grid-12col  [CONDENSED — 12 capabilities]')
label(doc, 'Section 14 — Training Management Platform Capabilities  [EXISTING — CONDENSED]')
section_heading(doc, 'Every Package Includes Access to the Edstellar Training Platform', level=2)

body(doc, 'All Edstellar corporate training packages come with full access to our training management platform. The platform handles scheduling, reporting, trainer management, and learning analytics in one place, eliminating the administrative overhead that inflates the true cost of corporate training programs.')

label(doc, '12 Platform Capabilities (4x3 grid, icon + title + 1 line each)')
capabilities = [
    ('User Management',          'Manage all platform users across business units and locations'),
    ('Course Discovery',         'Browse 2,000+ technical, management, behavioral, and compliance programs'),
    ('Training Plan Scheduling', 'Build and schedule training plans to address skill gaps on time'),
    ('Verified Trainer Network', 'Access carefully vetted, qualified, and experienced trainers on demand'),
    ('HRMS Integration',         'Sync with existing HR platforms to track and manage training progress'),
    ('Logistics and Operations', 'Smooth training delivery through dedicated coordination support'),
    ('Skill Gap Analysis',       'Use the Skill Matrix to identify employee capability gaps through reports'),
    ('Certificate Management',   'Issue certificates to all participants upon program completion'),
    ('Dedicated Resources',      'Access tools and templates that enhance your training quality'),
    ('Stellar AI Suggestions',   'Get AI-driven training recommendations based on job role and skill profile'),
    ('Analytics and Reports',    'Track training progress, completion rates, and program effectiveness'),
    ('Feedback Collection',      'Gather structured participant feedback to improve future training plans'),
]
for cap_name, cap_desc in capabilities:
    p = doc.add_paragraph(style='List Bullet')
    r1 = p.add_run(cap_name + ': ')
    r1.bold = True; r1.font.name = 'Calibri'; r1.font.size = Pt(10.5)
    r2 = p.add_run(cap_desc)
    r2.font.name = 'Calibri'; r2.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(3)

body(doc, 'The platform is included at no additional cost with every corporate training package tier. Enterprise and Custom package clients also receive advanced analytics dashboards and white-labeled certificate templates.')

divider(doc)


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 15 — FREE TRIAL / RISK REVERSAL CTA
# ══════════════════════════════════════════════════════════════════════════════
tag_bar(doc, 'edstellar-cta-banner-lime')
label(doc, 'Section 15 — Free Trial and Risk Reversal CTA  [NEW]')
section_heading(doc, 'Not Sure Which Package Is Right? Start with a Paid Trial Session', level=2)

body(doc, 'Choosing the right corporate training package is a significant decision. That is why every Edstellar package comes with a paid trial session option before you commit to a full plan.')
body(doc, 'The trial session gives your team a full 8-hour instructor-led training experience, a post-session feedback report, and a utilization recommendation from your assigned LSM. Most teams upgrade to a full corporate training package within two weeks of completing their trial.')

label(doc, 'CTA Buttons')
bullet(doc, 'Primary (navy, on lime background): Book a Trial Session')
bullet(doc, 'Secondary (outline): Talk to Our Training Experts')

label(doc, 'Supporting Trust Line (below CTAs)')
body(doc, 'No long-term commitment required for the trial. Pricing applies only to the training session itself.')

divider(doc)


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 16 — TRUST AND CREDENTIALS STRIP
# ══════════════════════════════════════════════════════════════════════════════
tag_bar(doc, 'CUSTOM CREDENTIALS STRIP  [NEW — white background, horizontal]')
label(doc, 'Section 16 — Trust and Credentials Strip  [NEW]')
section_heading(doc, 'Recognized Quality. Verified Security. Proven at Scale.', level=2)

body(doc, 'Edstellar holds dual ISO certification, ensuring that our corporate training delivery meets internationally recognized standards for quality management and data security.')

label(doc, '4-Item Credentials Row (horizontal, centered, equal spacing)')
bullet(doc, 'ISO 9001:2015 Certified: Our training delivery, platform operations, and client management processes meet the international standard for quality management systems. Every corporate training program you book is backed by a certified quality framework.')
bullet(doc, 'ISO 27001:2022 Certified: Your organizational data, employee records, and training information are protected under the international standard for information security management. Your data never leaves our secured environment.')
bullet(doc, 'Global Delivery: Training delivered in 60+ countries. Virtual and onsite sessions available across all time zones with multilingual trainer options in key markets.')
bullet(doc, 'Trusted Since [Year]: Thousands of instructor-led training sessions delivered to enterprise teams across IT, BFSI, manufacturing, healthcare, and professional services sectors globally.')

note_box(doc, 'Developer note: Replace [Year] with Edstellar founding year. Add Clutch or G2 rating badge if available. Consider adding a "As Featured In" media strip if press coverage exists.')
divider(doc)


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 17 — EXPANDED FAQ
# ══════════════════════════════════════════════════════════════════════════════
tag_bar(doc, 'edstellar-faq-section  [EXPANDED — 17 questions]')
label(doc, 'Section 17 — Frequently Asked Questions About Corporate Training Prices and Packages  [EXPANDED]')
section_heading(doc, 'Corporate Training Pricing FAQs', level=2)

body(doc, 'These are the most common questions buyers ask when evaluating corporate training prices and packages. Answers are structured to address both practical and strategic concerns.')

faqs = [
    (
        'How much does corporate training cost per employee?',
        'The industry average for corporate training cost per employee is $774 per learner per year (ATD 2024). This varies by company size: small organizations spend around $1,047 per employee, mid-size companies average $739, and large enterprises bring this down to $398 through economies of scale. Edstellar packages are designed to deliver competitive cost per learner especially for teams committing to annual training programs.'
    ),
    (
        'How are corporate training fees structured at Edstellar?',
        'Edstellar training fees are structured on a trainee license model. One license equals one 8-hour session for one participant. Packages are pre-purchased in four tiers: Starter (120 trainee slots), Growth (320 slots), Enterprise (800 slots), and Custom (unlimited). This model gives organizations full cost predictability and flexibility in how they allocate training hours across teams and programs.'
    ),
    (
        'What is included in Edstellar corporate training packages?',
        'Every package includes trainer fees, course materials, handouts, access to 2,000+ programs, virtual or onsite delivery, a Dedicated Learning Services Manager, retrospective sessions (2 hours per training program), and access to the Edstellar training management platform. Certification and exam fees, lab access, and post-training support are available separately.'
    ),
    (
        'What is the difference between the Starter, Growth, and Enterprise packages?',
        'The packages differ in the number of trainee license slots, total training hours, session count, and validity period. Starter is designed for teams of up to 15 with up to 64 hours of training over 6 months. Growth suits teams of 16 to 40 with up to 160 hours over 9 months and saves up to 10 percent compared to per-session rates. Enterprise covers teams of 40 to 100 with up to 400 hours over 12 months and saves up to 20 percent. Custom is for unlimited teams and unlimited sessions priced on a per-employee basis.'
    ),
    (
        'What are the charges for onsite corporate training?',
        'Onsite in-person training is charged at 30 percent above the standard package rate to cover trainer travel and accommodation expenses. This applies when training is delivered at your office, a client-preferred venue, or any location outside virtual delivery. Additional fees may apply for remote locations or international assignments. Contact our team for an onsite training quote specific to your geography.'
    ),
    (
        'Can I get training without purchasing a full package?',
        'Yes. Edstellar offers training options for individual employees and small groups without requiring a package commitment. A paid trial session is also available, giving your team a full 8-hour instructor-led training experience before you decide on a package. Contact us to discuss single-session or ad-hoc training pricing for your specific needs.'
    ),
    (
        'Are there multi-seat discounts for teams?',
        'Yes. The package pricing model itself functions as a volume discount. The larger the package tier, the lower the effective cost per trainee license. The Growth package saves up to 10 percent and the Enterprise package saves up to 20 percent compared to booking sessions individually. Custom packages for very large teams are priced on a per-employee basis for maximum cost efficiency at scale.'
    ),
    (
        'Can I mix virtual and onsite training within one package?',
        'Yes. All packages support both virtual instructor-led training (VILT) and onsite delivery. You can allocate different sessions within the same package to different delivery formats depending on the program requirements and team location. Onsite sessions incur the standard 30 percent delivery surcharge on top of the license cost.'
    ),
    (
        'What additional services can I add to a corporate training package?',
        'Additional services available as paid add-ons include executive coaching, individual mentoring, Training Needs Analysis (TNA), custom content development, and training delivered at an external venue. These services are coordinated through your Dedicated LSM and billed separately from the core package.'
    ),
    (
        'How long does it take to set up a training package?',
        'Most packages are activated within 5 to 7 business days of contract signing. Your Dedicated Learning Services Manager initiates the onboarding call within 24 hours of activation to begin calendar planning. First training sessions can typically be scheduled within two to three weeks of package activation, depending on trainer availability and your team scheduling.'
    ),
    (
        'What ROI can I expect from a corporate training package?',
        'ROI from corporate training depends on the programs selected and how outcomes are measured. Common indicators include reduced employee turnover (replacing one employee costs 50 to 200 percent of their annual salary), faster time to competency (30 to 50 percent faster than unstructured on-the-job learning), improved compliance audit scores, and measurable performance improvements in sales, operations, or technical output. Edstellar clients regularly report specific outcome gains within 3 to 6 months of program completion.'
    ),
    (
        'Do you offer corporate training packages for small teams?',
        'Yes. The Starter package is specifically designed for smaller organizations or teams with limited but structured training requirements. It provides 120 trainee slots and 64 hours of training over 6 months. For teams of fewer than 10 people, individual session booking or a paid trial may be more cost-effective. Speak to our team and we will recommend the most affordable training option for your team size.'
    ),
    (
        'What is the role of the Dedicated Learning Services Manager?',
        'The Dedicated Learning Services Manager (LSM) is your single point of contact for the entire engagement. The LSM handles onboarding, session scheduling, trainer coordination, pre-training briefing, logistics for onsite sessions, post-session feedback collection, retrospective facilitation, monthly utilization reporting, and renewal planning. They are available throughout your package validity period and ensure that your training program runs on schedule and to quality.'
    ),
    (
        'Do you offer corporate training in Dubai, India, the USA, and globally?',
        'Yes. Edstellar delivers corporate training globally across 60+ countries. Virtual instructor-led sessions are available in all time zones. Onsite training is available in all major markets including the USA, UK, India, UAE, Australia, Southeast Asia, and the Middle East. International onsite engagements are priced with a 30 percent delivery surcharge plus applicable travel costs.'
    ),
    (
        'How do Edstellar corporate training prices compare with other providers?',
        'Edstellar packages are competitively priced against both specialist training companies and eLearning platforms. The key differentiator is the combination of live expert trainers, a 2,000+ course library, a managed service layer (LSM), and a training platform included in one package price. Most comparable providers charge separately for each of these elements, making the total cost of working with ad-hoc vendors significantly higher than the headline per-session rate suggests.'
    ),
    (
        'Are Edstellar training packages available for IT and technical training?',
        'Yes. Edstellar has one of the broadest IT and technical training libraries available through a managed service model: 1,005+ IT and technical programs including cloud computing (105+), cybersecurity (30+), data analytics (50+), DevOps (35+), AI and machine learning (135+), project management (20+), and data science (15+). All programs are delivered by certified technical trainers with direct hands-on industry experience.'
    ),
    (
        'Can I track training progress and completion through the package?',
        'Yes. The Edstellar training platform provides full visibility into training progress, session completion rates, individual participant records, skill gap tracking, and program-level analytics. Monthly reports prepared by your LSM summarize utilization against your package allocation and flag any at-risk training milestones.'
    ),
]

for q, a in faqs:
    p = doc.add_paragraph()
    r = p.add_run('Q: ' + q)
    r.bold = True; r.font.name = 'Calibri'; r.font.size = Pt(10.5)
    r.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    body(doc, a)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)

note_box(doc, 'SEO note: FAQ section targets People Also Ask boxes for "how much does corporate training cost", "corporate training fees", "corporate training packages", "corporate training pricing models", "onsite training pricing", "multi-seat discounts", and 10+ additional high-impression queries from Search Console data.')
divider(doc)


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 18 — CONTACT AND QUOTE FORM
# ══════════════════════════════════════════════════════════════════════════════
tag_bar(doc, 'edstellar-form-section  [ENHANCED]')
label(doc, 'Section 18 — Contact and Quote Form  [ENHANCED]')
section_heading(doc, 'Get a Custom Corporate Training Quote for Your Organization', level=2)

body(doc, 'Tell us about your training requirements and we will build a custom package recommendation within one business day. No obligation. No commitment required.')

label(doc, 'Form Fields (8 fields minimum)')
bullet(doc, 'First Name (required)')
bullet(doc, 'Last Name (required)')
bullet(doc, 'Work Email (required)')
bullet(doc, 'Company Name (required)')
bullet(doc, 'Country (dropdown, required)')
bullet(doc, 'Phone Number with Country Code (required)')
bullet(doc, 'Approximate Team Size (dropdown: 1-15 / 16-50 / 51-200 / 201-500 / 500+)')
bullet(doc, 'Primary Training Goal (dropdown: Technical / Leadership / Compliance / Behavioral / Management / Other)')
bullet(doc, 'Message / Specific Requirements (optional text area)')

label(doc, 'Submit Button')
body(doc, 'Label: Get My Custom Quote')

label(doc, 'Trust Signals Near Form')
bullet(doc, 'We respond within 1 business day')
bullet(doc, 'Your data is protected under ISO 27001:2022 certified security standards')
bullet(doc, 'No long-term commitment required for initial consultation')

label(doc, 'Contact Details (shown alongside the form)')
bullet(doc, 'USA: +1 682 297 4830')
bullet(doc, 'United Kingdom: +44 151 453 4009')
bullet(doc, 'India: +91 636 483 3830')
bullet(doc, 'Email: contact@edstellar.com')

divider(doc)


# ══════════════════════════════════════════════════════════════════════════════
#  FOOTER NOTE
# ══════════════════════════════════════════════════════════════════════════════
tag_bar(doc, 'edstellar-footer  [EXISTING]')
label(doc, 'Section 19 — Footer  [EXISTING]')
body(doc, 'Use the standard Edstellar site footer with ISO certification badges, legal links, social media icons, and the full navigation link columns. No content changes required.')

doc.add_paragraph()
add_hr(doc, '2D2F6B')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Edstellar Corporate Training Pricing Page  |  Content Document  |  April 2026')
r.font.size = Pt(9); r.font.color.rgb = GREY; r.font.name = 'Calibri'


# ══════════════════════════════════════════════════════════════════════════════
#  SAVE
# ══════════════════════════════════════════════════════════════════════════════
out = r'C:\Users\Edstellar\Documents\Consulting Page Builder\Pricing\pricing-page-updated.docx'
doc.save(out)
print(f'Saved: {out}')
